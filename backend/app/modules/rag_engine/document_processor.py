import asyncio
import logging
from typing import Dict, Any, List, Optional
from pathlib import Path
from datetime import datetime

import PyPDF2
import docx
from openpyxl import load_workbook

from .vector_service import VectorRAGService

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """
    Processore documenti per RAG Engine
    Supporta: PDF, DOCX, XLSX, TXT, MD
    """
    
    def __init__(self):
        self.vector_service = VectorRAGService()
        self.supported_formats = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.xlsx': self._extract_xlsx,
            '.txt': self._extract_txt,
            '.md': self._extract_txt
        }
    
    async def _extract_pdf(self, file_path: Path) -> str:
        """
        Estrae testo da PDF
        """
        content = ""
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    content += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"Error extracting PDF {file_path}: {e}")
            raise
        
        return content.strip()
    
    async def _extract_docx(self, file_path: Path) -> str:
        """
        Estrae testo da DOCX
        """
        try:
            doc = docx.Document(file_path)
            content = []
            
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            
            return "\n".join(content).strip()
        except Exception as e:
            logger.error(f"Error extracting DOCX {file_path}: {e}")
            raise
    
    async def _extract_xlsx(self, file_path: Path) -> str:
        """
        Estrae testo da XLSX
        """
        try:
            workbook = load_workbook(file_path, read_only=True)
            content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                content.append(f"Sheet: {sheet_name}\n")
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = "\t".join([str(cell) if cell is not None else "" for cell in row])
                    if row_text.strip():
                        content.append(row_text)
                
                content.append("\n")
            
            return "\n".join(content).strip()
        except Exception as e:
            logger.error(f"Error extracting XLSX {file_path}: {e}")
            raise
    
    async def _extract_txt(self, file_path: Path) -> str:
        """
        Estrae testo da TXT/MD
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except Exception as e:
            logger.error(f"Error extracting TXT {file_path}: {e}")
            raise
    
    def get_supported_formats(self) -> List[str]:
        """
        Ritorna i formati supportati
        """
        return list(self.supported_formats.keys())
    
    def health_check(self) -> Dict[str, Any]:
        """
        Health check del document processor
        """
        return {
            'supported_formats': self.get_supported_formats(),
            'vector_service_health': self.vector_service.health_check()
        }

    async def extract_text(self, file_path):
        """
        Estrai testo da documento con metadata
        
        Returns:
            {
                'text': str,
                'metadata': dict,
                'success': bool,
                'error': str|None
            }
        """
        try:
            from pathlib import Path
            from datetime import datetime
            
            file_path = Path(file_path)
            
            if not file_path.exists():
                return {
                    'text': '',
                    'metadata': {},
                    'success': False,
                    'error': f'File not found: {file_path}'
                }
            
            file_extension = file_path.suffix.lower()
            
            if file_extension not in self.supported_formats:
                return {
                    'text': '',
                    'metadata': {},
                    'success': False,
                    'error': f'Unsupported format: {file_extension}'
                }
            
            # Metadata base
            metadata = {
                'filename': file_path.name,
                'file_size': file_path.stat().st_size,
                'format': file_extension,
                'processed_at': datetime.utcnow().isoformat()
            }
            
            # Estrazione testo per formato
            if file_extension == '.txt' or file_extension == '.md':
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            elif file_extension == '.pdf':
                # Estrazione PDF con PyPDF2
                import PyPDF2
                text = ""
                try:
                    with open(file_path, 'rb') as f:
                        pdf_reader = PyPDF2.PdfReader(f)
                        for page_num in range(min(len(pdf_reader.pages), 10)):  # Max 10 pagine
                            page = pdf_reader.pages[page_num]
                            text += page.extract_text() + "\n"
                except Exception as e:
                    text = f"Errore estrazione PDF: {str(e)}"
            elif file_extension == '.docx':
                # Estrazione DOCX con python-docx
                import docx
                text = ""
                try:
                    doc = docx.Document(file_path)
                    for paragraph in doc.paragraphs:
                        text += paragraph.text + "\n"
                    
                    # Estrai anche dalle tabelle
                    for table in doc.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                text += cell.text + " "
                            text += "\n"
                except Exception as e:
                    text = f"Errore estrazione DOCX: {str(e)}"
            else:
                # Per altri formati, usa una versione semplificata
                text = f"Documento {file_path.name} - Formato {file_extension} - Processing completo da implementare"
            
            return {
                'text': text,
                'metadata': metadata,
                'success': True,
                'error': None
            }
            
        except Exception as e:
            return {
                'text': '',
                'metadata': {},
                'success': False,
                'error': str(e)
            }
